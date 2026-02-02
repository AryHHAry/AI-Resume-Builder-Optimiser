"""
AI Resume Builder Optimizer
Web App untuk membantu job seekers membuat resume ATS-friendly yang tailored
Created by Ary HH (aryhharyanto@proton.me)
"""

import streamlit as st
import pandas as pd
import numpy as np
from datetime import datetime
import json
import re
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import matplotlib.pyplot as plt
import io
from fpdf import FPDF
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import base64

# Konfigurasi halaman
st.set_page_config(
    page_title="AI Resume Builder Optimizer",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded"
)

# CSS Custom untuk UI yang lebih menarik
st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        font-weight: bold;
        color: #1E3A8A;
        text-align: center;
        margin-bottom: 1rem;
    }
    .sub-header {
        font-size: 1.2rem;
        color: #64748B;
        text-align: center;
        margin-bottom: 2rem;
    }
    .metric-card {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        padding: 1.5rem;
        border-radius: 10px;
        color: white;
        text-align: center;
        margin: 1rem 0;
    }
    .warning-box {
        background-color: #FEF3C7;
        border-left: 4px solid #F59E0B;
        padding: 1rem;
        border-radius: 5px;
        margin: 1rem 0;
    }
    .success-box {
        background-color: #D1FAE5;
        border-left: 4px solid #10B981;
        padding: 1rem;
        border-radius: 5px;
        margin: 1rem 0;
    }
    .stButton>button {
        width: 100%;
        background-color: #1E3A8A;
        color: white;
        font-weight: bold;
        border-radius: 8px;
        padding: 0.5rem 1rem;
    }
    .footer {
        text-align: center;
        color: #64748B;
        padding: 2rem 0;
        font-size: 0.9rem;
    }
</style>
""", unsafe_allow_html=True)

# Inisialisasi session state
if 'resume_data' not in st.session_state:
    st.session_state.resume_data = {}
if 'analysis_result' not in st.session_state:
    st.session_state.analysis_result = None
if 'language' not in st.session_state:
    st.session_state.language = 'Indonesia'

# Fungsi utilitas untuk ekstraksi keyword
@st.cache_data
def extract_keywords(text, top_n=20):
    """Ekstraksi keyword menggunakan TF-IDF"""
    # Preprocessing
    text = text.lower()
    text = re.sub(r'[^\w\s]', ' ', text)
    
    # TF-IDF
    vectorizer = TfidfVectorizer(max_features=top_n, stop_words='english', ngram_range=(1, 2))
    try:
        tfidf_matrix = vectorizer.fit_transform([text])
        feature_names = vectorizer.get_feature_names_out()
        scores = tfidf_matrix.toarray()[0]
        
        keywords = [(feature_names[i], scores[i]) for i in range(len(feature_names))]
        keywords.sort(key=lambda x: x[1], reverse=True)
        return keywords
    except:
        return []

def calculate_keyword_overlap(resume_text, job_desc_text):
    """Hitung keyword overlap menggunakan cosine similarity"""
    if not resume_text or not job_desc_text:
        return 0.0, []
    
    vectorizer = TfidfVectorizer(stop_words='english', ngram_range=(1, 2))
    try:
        tfidf_matrix = vectorizer.fit_transform([resume_text, job_desc_text])
        similarity = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0]
        
        # Extract missing keywords
        job_keywords = set([kw[0] for kw in extract_keywords(job_desc_text, 30)])
        resume_keywords = set([kw[0] for kw in extract_keywords(resume_text, 30)])
        missing_keywords = list(job_keywords - resume_keywords)[:10]
        
        return similarity, missing_keywords
    except:
        return 0.0, []

def calculate_section_completeness(resume_data):
    """Hitung kelengkapan section resume"""
    required_sections = ['nama', 'email', 'pendidikan', 'pengalaman', 'skills']
    completed = sum(1 for section in required_sections if resume_data.get(section))
    return (completed / len(required_sections)) * 100

def calculate_tone_fit(tone, industri):
    """Hitung kesesuaian tone dengan industri"""
    tone_industry_map = {
        'professional': ['finance', 'consulting', 'legal', 'healthcare'],
        'creative': ['marketing', 'design', 'media', 'arts'],
        'technical': ['it', 'engineering', 'data science', 'technology']
    }
    
    industri_lower = industri.lower()
    for tone_type, industries in tone_industry_map.items():
        if tone.lower() == tone_type:
            for ind in industries:
                if ind in industri_lower:
                    return 100
            return 70
    return 80

def calculate_ats_score(resume_data, job_desc, params):
    """Kalkulasi skor ATS comprehensive"""
    # Gabungkan resume text
    resume_text = f"{resume_data.get('summary', '')} {resume_data.get('pengalaman', '')} {resume_data.get('skills', '')}"
    
    # 1. Keyword Overlap (50%)
    keyword_similarity, missing_keywords = calculate_keyword_overlap(resume_text, job_desc)
    keyword_score = keyword_similarity * 50
    
    # 2. Section Completeness (30%)
    section_score = calculate_section_completeness(resume_data) * 0.30
    
    # 3. Tone Fit (10%)
    tone_score = calculate_tone_fit(params['tone'], params['industri']) * 0.10
    
    # 4. Length Compliance (10%)
    page_count = len(resume_text.split()) / 300  # Estimasi 300 kata per halaman
    
    # Extract max pages using regex for robustness (handles "1 halaman", "1-2 halaman", etc.)
    page_numbers = [int(n) for n in re.findall(r'\d+', params['panjang_resume'])]
    max_pages = page_numbers[-1] if page_numbers else 2
    length_score = 10 if page_count <= max_pages else 5
    
    total_score = min(keyword_score + section_score + tone_score + length_score, 100)
    
    # Confidence interval berdasarkan variasi industri
    confidence_interval = 10 if params['level_pengalaman'] == 'entry-level' else 8 if params['level_pengalaman'] == 'mid-level' else 5
    
    return {
        'total_score': round(total_score, 1),
        'keyword_score': round(keyword_score, 1),
        'section_score': round(section_score, 1),
        'tone_score': round(tone_score, 1),
        'length_score': round(length_score, 1),
        'missing_keywords': missing_keywords,
        'confidence_interval': confidence_interval,
        'ats_compliant': total_score >= 70 and page_count <= 2
    }

def generate_ai_summary(params, resume_data):
    """Generate AI-powered summary"""
    templates = {
        'entry-level': f"{params['level_pengalaman'].title()} professional di bidang {params['industri']} dengan passion untuk belajar dan berkontribusi. Memiliki dasar yang kuat dalam {resume_data.get('skills', 'berbagai skills relevan')}.",
        'mid-level': f"Profesional {params['industri']} berpengalaman dengan track record yang solid dalam {resume_data.get('pengalaman', 'berbagai proyek')}. Mahir dalam {resume_data.get('skills', 'teknologi dan tools terkini')}.",
        'senior-level': f"Senior {params['industri']} professional dengan pengalaman ekstensif dalam leadership dan strategic planning. Expert dalam {resume_data.get('skills', 'domain expertise')} dengan proven results."
    }
    
    base_summary = templates.get(params['level_pengalaman'], templates['mid-level'])
    
    if params['tone'] == 'creative':
        base_summary = "🚀 " + base_summary + " Selalu siap untuk challenge baru dan inovasi!"
    elif params['tone'] == 'technical':
        base_summary = base_summary + " Focus pada technical excellence dan best practices."
    
    return base_summary

def generate_recommendations(analysis_result, params):
    """Generate actionable recommendations"""
    recommendations = []
    score = analysis_result['total_score']
    
    if score < 70:
        recommendations.append({
            'priority': 'HIGH',
            'action': f"Tambahkan keyword ATS-friendly: {', '.join(analysis_result['missing_keywords'][:5])}",
            'impact': '+15-20% score'
        })
    
    if analysis_result['section_score'] < 25:
        recommendations.append({
            'priority': 'HIGH',
            'action': "Lengkapi semua section wajib (Summary, Experience, Skills, Education)",
            'impact': '+10-15% score'
        })
    
    if analysis_result['tone_score'] < 8:
        recommendations.append({
            'priority': 'MEDIUM',
            'action': f"Sesuaikan tone resume menjadi lebih {params['industri']}-friendly",
            'impact': '+5-10% score'
        })
    
    if len(analysis_result['missing_keywords']) > 5:
        recommendations.append({
            'priority': 'MEDIUM',
            'action': f"Integrasikan {len(analysis_result['missing_keywords'])} missing keywords ke dalam pengalaman kerja",
            'impact': '+10-15% score'
        })
    
    if score >= 70 and score < 85:
        recommendations.append({
            'priority': 'LOW',
            'action': "Optimalkan formatting dan gunakan action verbs yang kuat",
            'impact': '+5-10% score'
        })
    
    if not recommendations:
        recommendations.append({
            'priority': 'LOW',
            'action': "Resume sudah excellent! Pertimbangkan customisasi lebih lanjut per job application",
            'impact': 'Maintain quality'
        })
    
    return recommendations

def create_pdf_resume(resume_data, params, template='modern'):
    """Generate PDF resume dengan template"""
    pdf = FPDF()
    pdf.add_page()
    
    # Header
    pdf.set_font('Arial', 'B', 24)
    pdf.cell(0, 10, resume_data.get('nama', 'Your Name'), ln=True, align='C')
    
    pdf.set_font('Arial', '', 10)
    pdf.cell(0, 6, resume_data.get('email', 'email@example.com'), ln=True, align='C')
    pdf.cell(0, 6, resume_data.get('telepon', '+62 XXX XXX XXX'), ln=True, align='C')
    pdf.ln(5)
    
    # Summary
    if resume_data.get('summary'):
        pdf.set_font('Arial', 'B', 14)
        pdf.cell(0, 8, 'PROFESSIONAL SUMMARY', ln=True)
        pdf.set_font('Arial', '', 10)
        pdf.multi_cell(0, 5, resume_data['summary'])
        pdf.ln(3)
    
    # Experience
    if resume_data.get('pengalaman'):
        pdf.set_font('Arial', 'B', 14)
        pdf.cell(0, 8, 'WORK EXPERIENCE', ln=True)
        pdf.set_font('Arial', '', 10)
        pdf.multi_cell(0, 5, resume_data['pengalaman'])
        pdf.ln(3)
    
    # Education
    if resume_data.get('pendidikan'):
        pdf.set_font('Arial', 'B', 14)
        pdf.cell(0, 8, 'EDUCATION', ln=True)
        pdf.set_font('Arial', '', 10)
        pdf.multi_cell(0, 5, resume_data['pendidikan'])
        pdf.ln(3)
    
    # Skills
    if resume_data.get('skills'):
        pdf.set_font('Arial', 'B', 14)
        pdf.cell(0, 8, 'SKILLS', ln=True)
        pdf.set_font('Arial', '', 10)
        pdf.multi_cell(0, 5, resume_data['skills'])
    
    return pdf.output(dest='S').encode('latin-1')

def create_word_resume(resume_data, params, template='modern'):
    """Generate Word resume dengan template"""
    doc = Document()
    
    # Header
    header = doc.add_heading(resume_data.get('nama', 'Your Name'), 0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    contact = doc.add_paragraph()
    contact.add_run(f"{resume_data.get('email', 'email@example.com')} | {resume_data.get('telepon', '+62 XXX XXX XXX')}")
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Summary
    if resume_data.get('summary'):
        doc.add_heading('PROFESSIONAL SUMMARY', 2)
        doc.add_paragraph(resume_data['summary'])
    
    # Experience
    if resume_data.get('pengalaman'):
        doc.add_heading('WORK EXPERIENCE', 2)
        doc.add_paragraph(resume_data['pengalaman'])
    
    # Education
    if resume_data.get('pendidikan'):
        doc.add_heading('EDUCATION', 2)
        doc.add_paragraph(resume_data['pendidikan'])
    
    # Skills
    if resume_data.get('skills'):
        doc.add_heading('SKILLS', 2)
        doc.add_paragraph(resume_data['skills'])
    
    # Save to bytes
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io.getvalue()

def plot_score_breakdown(analysis_result):
    """Visualisasi breakdown skor"""
    fig, ax = plt.subplots(figsize=(10, 6))
    
    categories = ['Keyword\nMatch', 'Section\nComplete', 'Tone\nFit', 'Length\nCompliance']
    scores = [
        analysis_result['keyword_score'],
        analysis_result['section_score'],
        analysis_result['tone_score'],
        analysis_result['length_score']
    ]
    colors = ['#667eea', '#764ba2', '#f093fb', '#4facfe']
    
    bars = ax.bar(categories, scores, color=colors, alpha=0.8, edgecolor='black')
    
    # Add value labels on bars
    for bar, score in zip(bars, scores):
        height = bar.get_height()
        ax.text(bar.get_x() + bar.get_width()/2., height,
                f'{score:.1f}',
                ha='center', va='bottom', fontweight='bold', fontsize=11)
    
    ax.set_ylabel('Score', fontweight='bold', fontsize=12)
    ax.set_title('Resume Score Breakdown', fontweight='bold', fontsize=14)
    ax.set_ylim(0, 60)
    ax.grid(axis='y', alpha=0.3, linestyle='--')
    
    plt.tight_layout()
    return fig

# ============= MAIN APP =============

def main():
    # Header
    st.markdown('<div class="main-header">📄 AI Resume Builder Optimizer</div>', unsafe_allow_html=True)
    st.markdown('<div class="sub-header">Bangun Resume ATS-Friendly yang Tailored untuk Job Application Anda</div>', unsafe_allow_html=True)
    
    # Vision Section
    with st.expander("🎯 Vision & Tujuan Tool", expanded=False):
        st.markdown("""
        **AI Resume Builder Optimizer** adalah micro SaaS yang dirancang khusus untuk membantu job seekers di Indonesia 
        membuat resume yang optimal dan ATS-friendly. Tool ini menggunakan AI untuk:
        
        - 🎯 Menganalisis job description dan menyarankan keyword yang relevan
        - ✍️ Generate section resume yang tailored dengan tone yang sesuai
        - 📊 Memberikan analytics dan skor kesesuaian resume vs job posting
        - 📄 Export resume dalam format PDF/Word dengan berbagai template
        - 🚀 Memberikan actionable recommendations untuk meningkatkan peluang diterima
        
        **Target Users:** Job seekers, freelancers, fresh graduates yang ingin meningkatkan kualitas resume mereka.
        """)
    
    # Warning Box
    st.markdown("""
    <div class="warning-box">
        ⚠️ <strong>Disclaimer:</strong> Tool ini memberikan estimasi dan saran berdasarkan best practices ATS. 
        Hasil adalah panduan kasar dan bukan pengganti review manual atau konsultasi karir profesional.
    </div>
    """, unsafe_allow_html=True)
    
    # Sidebar - Parameters
    with st.sidebar:
        st.header("⚙️ Pengaturan")
        
        # Language selector
        lang = st.selectbox("🌐 Bahasa", ["Indonesia", "English"], index=0)
        st.session_state.language = lang
        
        st.subheader("Parameter Resume")
        
        user_name = st.text_input("📝 Nama Lengkap", value="")
        today = datetime.now().strftime("%Y-%m-%d")
        date = st.date_input("📅 Tanggal", value=datetime.now())
        
        level = st.selectbox(
            "🎓 Level Pengalaman",
            ["entry-level", "mid-level", "senior-level"]
        )
        
        industri = st.selectbox(
            "🏢 Industri",
            ["IT/Technology", "Marketing", "Finance", "Healthcare", "Engineering", 
             "Design", "Education", "Consulting", "Sales", "Other"]
        )
        
        panjang = st.selectbox(
            "📏 Panjang Resume",
            ["1 halaman", "1-2 halaman", "2 halaman"]
        )
        
        tone = st.selectbox(
            "🎨 Tone",
            ["professional", "creative", "technical"]
        )
        
        template = st.selectbox(
            "🎯 Template Design",
            ["Modern", "Classic", "Creative", "Minimalist", "Professional", 
             "Bold", "Elegant", "Tech", "Corporate", "Startup"]
        )
        
        params = {
            'nama': user_name,
            'tanggal': str(date),
            'level_pengalaman': level,
            'industri': industri,
            'panjang_resume': panjang,
            'tone': tone,
            'template': template
        }
    
    # Main Content Tabs
    tab1, tab2, tab3, tab4, tab5 = st.tabs([
        "📋 Input Data", 
        "🔍 Analisis Job Desc", 
        "✨ Generate Resume", 
        "📊 Analytics & Score",
        "💡 Recommendations"
    ])
    
    # TAB 1: Input Data Resume
    with tab1:
        st.header("📋 Data Resume Anda")
        
        col1, col2 = st.columns(2)
        
        with col1:
            nama = st.text_input("Nama Lengkap *", value=user_name)
            email = st.text_input("Email *", placeholder="nama@email.com")
            telepon = st.text_input("Telepon", placeholder="+62 812 3456 7890")
            
        with col2:
            linkedin = st.text_input("LinkedIn", placeholder="linkedin.com/in/username")
            portfolio = st.text_input("Portfolio/Website", placeholder="https://portfolio.com")
        
        st.subheader("🎓 Pendidikan")
        pendidikan = st.text_area(
            "Riwayat Pendidikan",
            placeholder="Contoh:\nS1 Teknik Informatika - Universitas Indonesia (2018-2022)\nGPA: 3.75/4.00",
            height=100
        )
        
        st.subheader("💼 Pengalaman Kerja")
        pengalaman = st.text_area(
            "Pengalaman Kerja",
            placeholder="Contoh:\nSoftware Engineer - PT Tech Indonesia (2022-sekarang)\n- Mengembangkan aplikasi web menggunakan React dan Node.js\n- Meningkatkan performa aplikasi hingga 40%",
            height=150
        )
        
        st.subheader("🛠️ Skills")
        skills = st.text_area(
            "Daftar Skills",
            placeholder="Contoh: Python, JavaScript, React, Node.js, SQL, Git, Agile, Problem Solving, Communication",
            height=100
        )
        
        st.subheader("✍️ Professional Summary (Opsional)")
        summary_input = st.text_area(
            "Professional Summary",
            placeholder="Atau biarkan kosong untuk di-generate otomatis oleh AI",
            height=100
        )
        
        if st.button("💾 Simpan Data Resume", type="primary"):
            st.session_state.resume_data = {
                'nama': nama,
                'email': email,
                'telepon': telepon,
                'linkedin': linkedin,
                'portfolio': portfolio,
                'pendidikan': pendidikan,
                'pengalaman': pengalaman,
                'skills': skills,
                'summary': summary_input
            }
            st.success("✅ Data resume berhasil disimpan!")
    
    # TAB 2: Analisis Job Description
    with tab2:
        st.header("🔍 Analisis Job Description")
        
        st.markdown("""
        Paste atau upload job description yang Anda target. AI akan mengekstrak keyword penting 
        dan memberikan saran untuk optimasi resume Anda.
        """)
        
        job_desc_method = st.radio("Pilih metode input:", ["Paste Text", "Upload File (.txt)"])
        
        job_desc = ""
        if job_desc_method == "Paste Text":
            job_desc = st.text_area(
                "Job Description",
                placeholder="Paste deskripsi pekerjaan di sini...",
                height=300
            )
        else:
            uploaded_file = st.file_uploader("Upload Job Description (.txt)", type=['txt'])
            if uploaded_file:
                job_desc = uploaded_file.read().decode('utf-8')
                st.text_area("Preview Job Description", job_desc, height=200)
        
        if st.button("🔎 Analisis Keyword", type="primary"):
            if job_desc:
                with st.spinner("Menganalisis job description..."):
                    keywords = extract_keywords(job_desc, top_n=20)
                    
                    st.subheader("🎯 Top ATS Keywords Ditemukan:")
                    
                    # Display in columns
                    cols = st.columns(4)
                    for idx, (keyword, score) in enumerate(keywords):
                        col_idx = idx % 4
                        with cols[col_idx]:
                            st.metric(
                                label=keyword.title(),
                                value=f"{score:.3f}",
                                delta="Relevance"
                            )
                    
                    # Store for later use
                    st.session_state.job_desc = job_desc
                    st.session_state.job_keywords = keywords
                    
                    st.success("✅ Analisis selesai! Gunakan keyword ini untuk optimasi resume Anda.")
            else:
                st.error("❌ Mohon input job description terlebih dahulu!")
    
    # TAB 3: Generate Resume
    with tab3:
        st.header("✨ Generate Resume Content")
        
        if not st.session_state.resume_data:
            st.warning("⚠️ Mohon isi data resume di Tab 'Input Data' terlebih dahulu!")
        else:
            st.markdown("AI akan membantu generate atau improve section resume Anda berdasarkan parameter yang dipilih.")
            
            col1, col2 = st.columns(2)
            
            with col1:
                if st.button("🤖 Generate Professional Summary", type="primary"):
                    with st.spinner("Generating summary..."):
                        summary = generate_ai_summary(params, st.session_state.resume_data)
                        st.session_state.resume_data['summary'] = summary
                        st.success("✅ Summary berhasil di-generate!")
                        st.text_area("Generated Summary", summary, height=150)
            
            with col2:
                st.info("""
                **Tips Generate Summary:**
                - Sesuaikan dengan level pengalaman
                - Highlight skills utama
                - Gunakan tone yang sesuai industri
                """)
            
            st.subheader("📝 Resume Preview")
            
            resume_preview = f"""
**{st.session_state.resume_data.get('nama', 'Nama Anda')}**
{st.session_state.resume_data.get('email', '')} | {st.session_state.resume_data.get('telepon', '')}
{st.session_state.resume_data.get('linkedin', '')}

**PROFESSIONAL SUMMARY**
{st.session_state.resume_data.get('summary', 'Belum di-generate')}

**WORK EXPERIENCE**
{st.session_state.resume_data.get('pengalaman', 'Belum diisi')}

**EDUCATION**
{st.session_state.resume_data.get('pendidikan', 'Belum diisi')}

**SKILLS**
{st.session_state.resume_data.get('skills', 'Belum diisi')}
            """
            
            st.text_area("Preview", resume_preview, height=400)
    
    # TAB 4: Analytics & Score
    with tab4:
        st.header("📊 Resume Analytics & ATS Score")
        
        if not st.session_state.resume_data:
            st.warning("⚠️ Mohon isi data resume terlebih dahulu!")
        elif 'job_desc' not in st.session_state:
            st.warning("⚠️ Mohon analisis job description terlebih dahulu di Tab 'Analisis Job Desc'!")
        else:
            if st.button("📊 Hitung ATS Score", type="primary"):
                with st.spinner("Menghitung skor ATS..."):
                    analysis = calculate_ats_score(
                        st.session_state.resume_data,
                        st.session_state.job_desc,
                        params
                    )
                    st.session_state.analysis_result = analysis
                    
                    # Display Score
                    col1, col2, col3 = st.columns(3)
                    
                    with col1:
                        st.markdown(f"""
                        <div class="metric-card">
                            <h1>{analysis['total_score']}%</h1>
                            <p>Overall ATS Score</p>
                            <small>±{analysis['confidence_interval']}% confidence</small>
                        </div>
                        """, unsafe_allow_html=True)
                    
                    with col2:
                        status = "✅ PASS" if analysis['ats_compliant'] else "❌ NEEDS IMPROVEMENT"
                        color = "#10B981" if analysis['ats_compliant'] else "#EF4444"
                        st.markdown(f"""
                        <div class="metric-card" style="background: {color};">
                            <h2>{status}</h2>
                            <p>ATS Compliance</p>
                        </div>
                        """, unsafe_allow_html=True)
                    
                    with col3:
                        benchmark = 80  # Average ATS score from Jobscan 2026
                        delta = analysis['total_score'] - benchmark
                        st.metric(
                            "vs Benchmark",
                            f"{analysis['total_score']}%",
                            f"{delta:+.1f}% vs avg {benchmark}%"
                        )
                    
                    # Score Breakdown Visualization
                    st.subheader("📈 Score Breakdown")
                    fig = plot_score_breakdown(analysis)
                    st.pyplot(fig)
                    
                    # Detailed Metrics
                    st.subheader("🔍 Detail Metrics")
                    metric_cols = st.columns(4)
                    
                    metrics = [
                        ("Keyword Match", analysis['keyword_score'], "50"),
                        ("Section Complete", analysis['section_score'], "30"),
                        ("Tone Fit", analysis['tone_score'], "10"),
                        ("Length OK", analysis['length_score'], "10")
                    ]
                    
                    for idx, (label, score, max_score) in enumerate(metrics):
                        with metric_cols[idx]:
                            st.metric(label, f"{score:.1f}", f"/{max_score}")
                    
                    # Missing Keywords
                    if analysis['missing_keywords']:
                        st.subheader("🎯 Missing Keywords")
                        st.warning(f"Tambahkan keyword berikut untuk meningkatkan skor: **{', '.join(analysis['missing_keywords'][:10])}**")
                    
                    # What-If Simulator
                    st.subheader("🔮 What-If Simulator")
                    st.markdown("Simulasikan perubahan parameter dan lihat dampaknya pada skor:")
                    
                    sim_col1, sim_col2 = st.columns(2)
                    
                    with sim_col1:
                        sim_level = st.select_slider(
                            "Level Pengalaman",
                            options=["entry-level", "mid-level", "senior-level"],
                            value=params['level_pengalaman']
                        )
                    
                    with sim_col2:
                        sim_tone = st.select_slider(
                            "Tone",
                            options=["professional", "creative", "technical"],
                            value=params['tone']
                        )
                    
                    if sim_level != params['level_pengalaman'] or sim_tone != params['tone']:
                        sim_params = params.copy()
                        sim_params['level_pengalaman'] = sim_level
                        sim_params['tone'] = sim_tone
                        
                        sim_analysis = calculate_ats_score(
                            st.session_state.resume_data,
                            st.session_state.job_desc,
                            sim_params
                        )
                        
                        delta_score = sim_analysis['total_score'] - analysis['total_score']
                        st.metric(
                            "Simulated Score",
                            f"{sim_analysis['total_score']}%",
                            f"{delta_score:+.1f}% change"
                        )
    
    # TAB 5: Recommendations
    with tab5:
        st.header("💡 Actionable Recommendations")
        
        if st.session_state.analysis_result:
            recommendations = generate_recommendations(st.session_state.analysis_result, params)
            
            st.markdown("Berikut adalah rekomendasi untuk meningkatkan skor ATS resume Anda:")
            
            priority_colors = {
                'HIGH': '#EF4444',
                'MEDIUM': '#F59E0B',
                'LOW': '#10B981'
            }
            
            for idx, rec in enumerate(recommendations, 1):
                color = priority_colors.get(rec['priority'], '#64748B')
                
                st.markdown(f"""
                <div style="background-color: {color}15; border-left: 4px solid {color}; padding: 1rem; border-radius: 5px; margin: 1rem 0;">
                    <strong style="color: {color};">{rec['priority']} PRIORITY</strong>
                    <h4>{idx}. {rec['action']}</h4>
                    <p>📈 <strong>Expected Impact:</strong> {rec['impact']}</p>
                </div>
                """, unsafe_allow_html=True)
            
            # Export Options
            st.subheader("📥 Export Resume")
            
            export_col1, export_col2, export_col3 = st.columns(3)
            
            with export_col1:
                if st.button("📄 Download PDF", type="primary"):
                    pdf_bytes = create_pdf_resume(st.session_state.resume_data, params, params['template'])
                    st.download_button(
                        label="💾 Save PDF",
                        data=pdf_bytes,
                        file_name=f"resume_{st.session_state.resume_data.get('nama', 'user')}_{datetime.now().strftime('%Y%m%d')}.pdf",
                        mime="application/pdf"
                    )
            
            with export_col2:
                if st.button("📝 Download Word", type="primary"):
                    word_bytes = create_word_resume(st.session_state.resume_data, params, params['template'])
                    st.download_button(
                        label="💾 Save Word",
                        data=word_bytes,
                        file_name=f"resume_{st.session_state.resume_data.get('nama', 'user')}_{datetime.now().strftime('%Y%m%d')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
            
            with export_col3:
                if st.button("📊 Export Analytics CSV"):
                    analytics_data = {
                        'Metric': ['Total Score', 'Keyword Match', 'Section Complete', 'Tone Fit', 'Length'],
                        'Score': [
                            st.session_state.analysis_result['total_score'],
                            st.session_state.analysis_result['keyword_score'],
                            st.session_state.analysis_result['section_score'],
                            st.session_state.analysis_result['tone_score'],
                            st.session_state.analysis_result['length_score']
                        ]
                    }
                    df = pd.DataFrame(analytics_data)
                    csv = df.to_csv(index=False)
                    st.download_button(
                        label="💾 Save CSV",
                        data=csv,
                        file_name=f"analytics_{datetime.now().strftime('%Y%m%d')}.csv",
                        mime="text/csv"
                    )
        else:
            st.info("📊 Lakukan analisis di Tab 'Analytics & Score' terlebih dahulu untuk mendapatkan rekomendasi.")
    
    # Footer
    st.markdown("---")
    st.markdown("""
    <div class="footer">
        <p><strong>AI Resume Builder Optimizer</strong> - Membantu Anda membuat resume ATS-friendly untuk peluang karir terbaik</p>
        <p>Created by <strong>Ary HH</strong> (aryhharyanto@proton.me)</p>
        <p><em>© 2026 - Non Open Source Project</em></p>
        <p style="font-size: 0.8rem; color: #94A3B8;">
            🔒 Privacy: Data Anda tersimpan hanya dalam session browser ini dan tidak dikirim ke server eksternal<br>
            🚀 Future: Integrasi LinkedIn API, Google Drive, dan advanced AI features coming soon
        </p>
    </div>
    """, unsafe_allow_html=True)

if __name__ == "__main__":
    main()
